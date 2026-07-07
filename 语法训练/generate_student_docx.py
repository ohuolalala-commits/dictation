# -*- coding: utf-8 -*-
"""生成 ASD 青少年英语语法训练·学生用书（知识讲解 + 练习题）"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime, shutil

doc = Document()

# ============ 页面设置 ============
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ============ 样式 ============
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
pf = style.paragraph_format
pf.line_spacing = 1.4
pf.space_after = Pt(4)

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = '微软雅黑'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if i == 1:
        h.font.size = Pt(20)
        h.font.color.rgb = RGBColor(0x4A, 0x90, 0xD9)
        h.font.bold = True
    elif i == 2:
        h.font.size = Pt(15)
        h.font.color.rgb = RGBColor(0x5C, 0xB8, 0x5C)
        h.font.bold = True

# ============ 颜色常量 ============
BLUE = (0x4A, 0x90, 0xD9)     # 主语
GREEN = (0x3D, 0x8B, 0x3D)    # Be动词
ORANGE = (0xE0, 0x8E, 0x2A)   # 动作动词
PURPLE = (0x8E, 0x44, 0xAD)   # 功能词
YELLOW_BG = (0xFF, 0xF3, 0xCD) # 黄色背景
GRAY = (0x88, 0x88, 0x88)

# ============ 辅助函数 ============
def P(text, bold=False, size=11, color=None, align=None, indent=None, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    if align is not None: p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: run.bold = True
    if color: run.font.color.rgb = RGBColor(*color)
    return p

def ColorWord(text, color):
    """在段落中插入彩色文字"""
    run = doc.add_paragraph().add_run(text)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(*color)
    run.bold = True
    return run

def Example(text, size=12, bold=False, indent=1.0, color=None):
    """例句，大号英文"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Georgia'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: run.bold = True
    if color: run.font.color.rgb = RGBColor(*color)
    return p

def Tip(text):
    """浅色提示"""
    P('💡 ' + text, size=9.5, color=GRAY, indent=0.8, space_before=2, space_after=2)

def Divider():
    P('─' * 50, size=8, color=(0xCC, 0xCC, 0xCC), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

def add_table(headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ============ 练习题辅助函数 ============
def ExTitle(num, title):
    P(f'练习 {num}：{title}', bold=True, size=11.5, space_before=10, space_after=6)

def FillBlank(items):
    """填空题: [("I ___ a boy.", ["am","is"], "am"), ...]"""
    for sentence, options, answer in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.2)
        p.paragraph_format.space_after = Pt(6)
        # 显示句子，空位用下划线
        run = p.add_run(sentence.replace('___', '_____'))
        run.font.size = Pt(12)
        run.font.name = 'Georgia'
        # 选项放在括号里
        opts = ' / '.join(options)
        run2 = p.add_run(f'    ( {opts} )')
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(*GRAY)

def ChooseCorrect(items):
    """选择题: ("He ___ a boy.", "A. am  B. is  C. are", "B")"""
    for i, (sentence, choices, answer) in enumerate(items, 1):
        P(f'{i}. {sentence}', size=11.5, indent=1.2, space_after=2)
        P(choices, size=10.5, indent=1.8, color=GRAY, space_after=6)

def BuildSentence(items):
    """组句题: 给乱序词，组正确句子"""
    for i, (words, answer) in enumerate(items, 1):
        P(f'{i}. 请把下面的词排成正确的句子：', size=10.5, indent=1.2, space_after=2)
        # 显示乱序词
        shuffled = words[:]
        import random
        random.shuffle(shuffled)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.8)
        p.paragraph_format.space_after = Pt(4)
        for w in shuffled:
            run = p.add_run(f'[ {w} ]  ')
            run.font.size = Pt(12)
            run.font.name = 'Georgia'
        # 答题横线
        P('答案：________________________________', size=11, indent=1.2, color=GRAY, space_after=8)

def WriteFromPic(items):
    """看图写句"""
    for i, (pic, hint) in enumerate(items, 1):
        P(f'{i}. {pic}', size=14, indent=1.2, space_after=2)
        if hint:
            P(f'   提示词：{hint}', size=9, color=GRAY, indent=1.8, space_after=2)
        P('   写句子：________________________________', size=11, indent=1.2, color=GRAY, space_after=8)

# ====================================================================
# 封面
# ====================================================================
for _ in range(5):
    doc.add_paragraph()

P('🧩', size=40, align=WD_ALIGN_PARAGRAPH.CENTER)
P('英语语法练习册', size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE)
P('English Grammar Workbook', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
doc.add_paragraph()
P('适合需要循序渐进的学习者', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
P('8 个单元 · 24 课 · 逐步掌握基础语法', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
doc.add_paragraph()
P(f'{datetime.date.today().strftime("%Y年%m月")}', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)

doc.add_page_break()

# ====================================================================
# 第0课：认识颜色标记
# ====================================================================
doc.add_heading('课前准备：认识颜色标记', level=1)

P('在这本练习册里，我们会用颜色来帮助你"看见"英语句子里的每个部分。就像拼图一样，每种颜色代表一种拼图块。', size=11)

P('')
P('🔵 蓝色 = "谁"（主语）', size=12, bold=True, color=BLUE, indent=1.0)
P('   比如：I（我）、He（他）、She（她）、a cat（一只猫）', size=10.5, indent=1.0, color=GRAY)

P('🟢 绿色 = "连接词"（Be 动词）', size=12, bold=True, color=GREEN, indent=1.0)
P('   比如：am、is、are', size=10.5, indent=1.0, color=GRAY)

P('🟠 橙色 = "动作"（动词）', size=12, bold=True, color=ORANGE, indent=1.0)
P('   比如：eat（吃）、run（跑）、play（玩）', size=10.5, indent=1.0, color=GRAY)

P('🟣 紫色 = "小帮手词"（功能词）', size=12, bold=True, color=PURPLE, indent=1.0)
P('   比如：a、the、can、in、on', size=10.5, indent=1.0, color=GRAY)

P('')
P('注意：我们不会同时用所有颜色！每次只学一种新的颜色。准备好了吗？', size=11, bold=True, space_before=8)

doc.add_page_break()

# ====================================================================
# 第1课：I am ...（我是...）
# ====================================================================
doc.add_heading('第一课  I am ...（我是...）', level=1)

# --- 知识讲解 ---
doc.add_heading('📖 知识讲解', level=2)

P('当我们想说"我是谁"或者"我感觉怎么样"的时候，用 I am。', size=11)

Example('I am a boy.', size=13)
P('（我是一个男孩。）', size=10, color=GRAY, indent=1.0)

Example('I am happy.', size=13)
P('（我很开心。）', size=10, color=GRAY, indent=1.0)

Example('I am tall.', size=13)
P('（我很高。）', size=10, color=GRAY, indent=1.0)

P('看出来了吗？句子= I + am + ____', bold=True, size=11, space_before=6)

doc.add_heading('🎨 颜色标记', level=2)
P('我们用颜色来看清楚句子的结构：', size=10.5, indent=0.5)

# 颜色示例
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
p.paragraph_format.space_before = Pt(8)
r1 = p.add_run('I ')
r1.font.size = Pt(14); r1.font.name = 'Georgia'; r1.font.color.rgb = RGBColor(*BLUE); r1.bold = True
r2 = p.add_run('am ')
r2.font.size = Pt(14); r2.font.name = 'Georgia'; r2.font.color.rgb = RGBColor(*GREEN); r2.bold = True
r3 = p.add_run('a boy.')
r3.font.size = Pt(14); r3.font.name = 'Georgia'; r3.font.color.rgb = RGBColor(*BLUE)

P('    ↑蓝色="我"    ↑绿色=连接词    ↑蓝色=男孩', size=9, color=GRAY, indent=1.0)

p2 = doc.add_paragraph()
p2.paragraph_format.left_indent = Cm(1.0)
r4 = p2.add_run('I ')
r4.font.size = Pt(14); r4.font.name = 'Georgia'; r4.font.color.rgb = RGBColor(*BLUE); r4.bold = True
r5 = p2.add_run('am ')
r5.font.size = Pt(14); r5.font.name = 'Georgia'; r5.font.color.rgb = RGBColor(*GREEN); r5.bold = True
r6 = p2.add_run('happy.')
r6.font.size = Pt(14); r6.font.name = 'Georgia'; r6.font.color.rgb = RGBColor(*BLUE)

Tip('记住：说"我"的时候，后面一定是 am。I am，I am，I am！')

Divider()

# --- 词汇表 ---
doc.add_heading('📝 本课单词', level=2)

vocab1 = [
    ['a boy', '男孩', 'a girl', '女孩'],
    ['a student', '学生', 'a teacher', '老师'],
    ['happy', '开心的', 'sad', '难过的'],
    ['tall', '高的', 'short', '矮的'],
    ['hungry', '饿的', 'tired', '累的'],
]
for row in vocab1:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(2)
    for j, word in enumerate(row):
        if j % 2 == 0:
            r = p.add_run(f'{word}  ')
            r.font.size = Pt(11); r.font.name = 'Georgia'; r.font.bold = True
        else:
            r = p.add_run(f'({word})    ')
            r.font.size = Pt(10); r.font.color.rgb = RGBColor(*GRAY)

Divider()

# --- 练习题 ---
doc.add_heading('✏️ 练习题', level=2)

ExTitle(1, '选择题（圈出正确的答案）')
ChooseCorrect([
    ('I ___ a boy.', 'A. am   B. is', 'A'),
    ('I ___ happy.', 'A. is   B. am', 'B'),
    ('I ___ a student.', 'A. are   B. am', 'B'),
    ('I ___ sad.', 'A. am   B. is', 'A'),
])

ExTitle(2, '填空题（选择 am 或 is 填进去）')
FillBlank([
    ('I ___ a girl.', ['am', 'is'], 'am'),
    ('I ___ tall.', ['am', 'is'], 'am'),
    ('I ___ hungry.', ['am', 'is'], 'am'),
    ('I ___ a teacher.', ['am', 'is'], 'am'),
])

ExTitle(3, '组句练习（把词排成正确的句子）')
BuildSentence([
    (['I', 'am', 'a', 'boy'], 'I am a boy.'),
    (['happy', 'I', 'am'], 'I am happy.'),
    (['tall', 'I', 'am'], 'I am tall.'),
    (['am', 'hungry', 'I'], 'I am hungry.'),
])

ExTitle(4, '看图写句子')
WriteFromPic([
    ('🧒（男孩）', 'I / am / a / boy'),
    ('😊（笑脸）', 'I / am / happy'),
    ('😢（难过）', ''),
    ('📏（高个子）', ''),
])

doc.add_page_break()

# ====================================================================
# 第2课：He is / She is（他是/她是）
# ====================================================================
doc.add_heading('第二课  He is / She is（他是... / 她是...）', level=1)

doc.add_heading('📖 知识讲解', level=2)

P('上一课我们学了说"我"的时候用 I am。现在来学说"他"和"她"。', size=11)

P('重要变化：', bold=True, size=11, space_before=6)
P('说"我"→ 用 am', size=11, indent=1.0)
P('说"他"→ 用 is', size=11, indent=1.0)
P('说"她"→ 也用 is', size=11, indent=1.0)

# 对比展示
P('', space_after=2)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
r1 = p.add_run('I am a boy.  ')
r1.font.size = Pt(12); r1.font.name = 'Georgia'
r2 = p.add_run('→  ')
r2.font.color.rgb = RGBColor(*GRAY)
r3 = p.add_run('He is a boy.')
r3.font.size = Pt(12); r3.font.name = 'Georgia'; r3.font.color.rgb = RGBColor(*GREEN)

P('👀 看到了吗？am 变成了 is！', bold=True, size=11, indent=1.0)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
r1 = p.add_run('I am a girl.  ')
r1.font.size = Pt(12); r1.font.name = 'Georgia'
r2 = p.add_run('→  ')
r2.font.color.rgb = RGBColor(*GRAY)
r3 = p.add_run('She is a girl.')
r3.font.size = Pt(12); r3.font.name = 'Georgia'; r3.font.color.rgb = RGBColor(*GREEN)

doc.add_heading('🎨 颜色标记', level=2)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.space_before = Pt(6)
r1 = p.add_run('He ')
r1.font.size = Pt(14); r1.font.name = 'Georgia'; r1.font.color.rgb = RGBColor(*BLUE); r1.bold = True
r2 = p.add_run('is ')
r2.font.size = Pt(14); r2.font.name = 'Georgia'; r2.font.color.rgb = RGBColor(*GREEN); r2.bold = True
r3 = p.add_run('tall.')
r3.font.size = Pt(14); r3.font.name = 'Georgia'; r3.font.color.rgb = RGBColor(*BLUE)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
r1 = p.add_run('She ')
r1.font.size = Pt(14); r1.font.name = 'Georgia'; r1.font.color.rgb = RGBColor(*BLUE); r1.bold = True
r2 = p.add_run('is ')
r2.font.size = Pt(14); r2.font.name = 'Georgia'; r2.font.color.rgb = RGBColor(*GREEN); r2.bold = True
r3 = p.add_run('happy.')
r3.font.size = Pt(14); r3.font.name = 'Georgia'; r3.font.color.rgb = RGBColor(*BLUE)

Tip('规则很简单：He 和 She 后面都用 is。记住：he is, she is, he is, she is！')

Divider()

doc.add_heading('✏️ 练习题', level=2)

ExTitle(1, '选择题')
ChooseCorrect([
    ('He ___ a boy.', 'A. am   B. is', 'B'),
    ('She ___ a girl.', 'A. am   B. is', 'B'),
    ('He ___ tall.', 'A. is   B. am', 'A'),
    ('She ___ happy.', 'A. am   B. is', 'B'),
])

ExTitle(2, '填空题（选 am 还是 is？）')
FillBlank([
    ('He ___ a boy.', ['am', 'is'], 'is'),
    ('She ___ happy.', ['am', 'is'], 'is'),
    ('I ___ a student.', ['am', 'is'], 'am'),
    ('He ___ tall.', ['am', 'is'], 'is'),
    ('I ___ hungry.', ['am', 'is'], 'am'),
    ('She ___ a teacher.', ['am', 'is'], 'is'),
])
Tip('第3题和第5题是复习哦！I 后面还是用 am，没有变！')

ExTitle(3, '组句练习')
BuildSentence([
    (['He', 'is', 'a', 'boy'], 'He is a boy.'),
    (['She', 'is', 'happy'], 'She is happy.'),
    (['is', 'He', 'tall'], 'He is tall.'),
])

ExTitle(4, '看图写句子')
WriteFromPic([
    ('👦（男孩）', 'He / is / a / boy'),
    ('👧（女孩）', 'She / is / a / girl'),
    ('👦😊（男孩+笑脸）', ''),
    ('👧📏（女孩+高）', ''),
])

doc.add_page_break()

# ====================================================================
# 第3课：You are / We are / They are
# ====================================================================
doc.add_heading('第三课  You are / We are / They are（你/我们/他们）', level=1)

doc.add_heading('📖 知识讲解', level=2)

P('我们已经学了：I → am，He/She → is。', size=11)
P('现在学新的：You、We、They 后面都用 are。', size=11, space_before=4)

# 对照表
P('人称和Be动词对照表：', bold=True, size=11, space_before=8)
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'
data = [
    ['I', 'am   ✓ 已学'],
    ['He / She', 'is   ✓ 已学'],
    ['You（你/你们）', 'are   🆕 新学'],
    ['We（我们）', 'are   🆕 新学'],
    ['They（他们/她们/它们）', 'are   🆕 新学'],
]
for i, (subj, verb) in enumerate(data):
    table.rows[i].cells[0].text = subj
    table.rows[i].cells[1].text = verb
    for cell in table.rows[i].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

P('')

Example('You are students.', size=12)
P('（你们是学生。）', size=10, color=GRAY, indent=1.0)

Example('We are happy.', size=12)
P('（我们很开心。）', size=10, color=GRAY, indent=1.0)

Example('They are tall.', size=12)
P('（他们很高。）', size=10, color=GRAY, indent=1.0)

Tip('You、We、They 这三个词后面都用 are。你可以这样记：除了 I 和 He/She，其他的都用 are！')

Divider()

doc.add_heading('✏️ 练习题', level=2)

ExTitle(1, '选择题（三选一：am / is / are）')
ChooseCorrect([
    ('You ___ students.', 'A. am   B. is   C. are', 'C'),
    ('We ___ happy.', 'A. am   B. is   C. are', 'C'),
    ('They ___ tall.', 'A. am   B. is   C. are', 'C'),
    ('I ___ a boy.', 'A. am   B. is   C. are', 'A'),
    ('He ___ tall.', 'A. am   B. is   C. are', 'B'),
    ('She ___ a girl.', 'A. am   B. is   C. are', 'B'),
])
Tip('最后三道是复习题！前面学的也不要忘哦。')

ExTitle(2, '填空题（am / is / are）')
FillBlank([
    ('You ___ a student.', ['am', 'is', 'are'], 'are'),
    ('We ___ happy.', ['am', 'is', 'are'], 'are'),
    ('They ___ tall.', ['am', 'is', 'are'], 'are'),
    ('I ___ hungry.', ['am', 'is', 'are'], 'am'),
    ('He ___ a teacher.', ['am', 'is', 'are'], 'is'),
    ('She ___ sad.', ['am', 'is', 'are'], 'is'),
])

ExTitle(3, '组句练习')
BuildSentence([
    (['You', 'are', 'students'], 'You are students.'),
    (['We', 'are', 'happy'], 'We are happy.'),
    (['are', 'They', 'tall'], 'They are tall.'),
    (['She', 'is', 'a', 'girl'], 'She is a girl.'),
])

ExTitle(4, '看图写句子')
WriteFromPic([
    ('👥（两个人+指向你）', 'You / are / students'),
    ('👥（一群人+自己）', 'We / are / happy'),
    ('👥（很多人）', ''),
    ('👧（女孩）', ''),
])

doc.add_page_break()

# ====================================================================
# 第4课：a 和 an（一个）
# ====================================================================
doc.add_heading('第四课  a 和 an（一个...）', level=1)

doc.add_heading('📖 知识讲解', level=2)

P('英语里，说"一个什么东西"的时候，名词前面要加 a 或 an。', size=11)

P('什么时候用 a？', bold=True, size=11, space_before=6)
P('大多数词的前面用 a：', size=10.5, indent=1.0)
Example('a cat （一只猫）', size=11, indent=2.0)
Example('a dog （一只狗）', size=11, indent=2.0)
Example('a book（一本书）', size=11, indent=2.0)

P('什么时候用 an？', bold=True, size=11, space_before=6)
P('如果单词的开头是 a, e, i, o, u 这几个字母，用 an：', size=10.5, indent=1.0)
Example('an apple （一个苹果）  ← apple 开头是 a', size=11, indent=2.0)
Example('an egg  （一个鸡蛋）  ← egg 开头是 e', size=11, indent=2.0)
Example('an umbrella（一把伞）← umbrella 开头是 u', size=11, indent=2.0)

Tip('小窍门：不用死记！看看单词的第一个字母。如果是 a/e/i/o/u → 用 an，其他 → 用 a。')

doc.add_heading('🎨 颜色标记', level=2)
P('a 和 an 都是"小帮手词"，用紫色：', size=10.5)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
r1 = p.add_run('a '); r1.font.size = Pt(14); r1.font.name = 'Georgia'; r1.font.color.rgb = RGBColor(*PURPLE); r1.bold = True
r2 = p.add_run('cat'); r2.font.size = Pt(14); r2.font.name = 'Georgia'; r2.font.color.rgb = RGBColor(*BLUE)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
r1 = p.add_run('an '); r1.font.size = Pt(14); r1.font.name = 'Georgia'; r1.font.color.rgb = RGBColor(*PURPLE); r1.bold = True
r2 = p.add_run('apple'); r2.font.size = Pt(14); r2.font.name = 'Georgia'; r2.font.color.rgb = RGBColor(*BLUE)

Divider()

doc.add_heading('✏️ 练习题', level=2)

ExTitle(1, '选择题（选 a 或 an）')
ChooseCorrect([
    ('___ cat', 'A. a   B. an', 'A'),
    ('___ apple', 'A. a   B. an', 'B'),
    ('___ dog', 'A. a   B. an', 'A'),
    ('___ egg', 'A. a   B. an', 'B'),
    ('___ umbrella', 'A. a   B. an', 'B'),
    ('___ book', 'A. a   B. an', 'A'),
    ('___ boy', 'A. a   B. an', 'A'),
    ('___ orange', 'A. a   B. an', 'B'),
])

ExTitle(2, '填写 a 或 an')
FillBlank([
    ('___ bird（鸟）', ['a', 'an'], 'a'),
    ('___ apple（苹果）', ['a', 'an'], 'an'),
    ('___ desk（桌子）', ['a', 'an'], 'a'),
    ('___ egg（鸡蛋）', ['a', 'an'], 'an'),
    ('___ girl（女孩）', ['a', 'an'], 'a'),
    ('___ elephant（大象）', ['a', 'an'], 'an'),
])

ExTitle(3, '看图写 a / an + 单词')
WriteFromPic([
    ('🐱', 'a / cat'),
    ('🍎', 'an / apple'),
    ('🐕', ''),
    ('🥚', ''),
    ('📚', ''),
    ('☂️', ''),
])

doc.add_page_break()

# ====================================================================
# 第5课：动作词（一般现在时）
# ====================================================================
doc.add_heading('第五课  动作词（I eat / You run / We play...）', level=1)

doc.add_heading('📖 知识讲解', level=2)

P('现在我们开始学"做动作"的句子。先说最简单的：I / You / We / They 做动作时，动词不变化。', size=11)

Example('I eat breakfast.  （我吃早饭。）', size=12)
Example('You go to school. （你去上学。）', size=12)
Example('We play soccer.   （我们踢足球。）', size=12)
Example('They read books.  （他们读书。）', size=12)

P('注意：这些句子里的动词都没有变！eat 就是 eat，go 就是 go。', bold=True, size=11, space_before=6)

doc.add_heading('🎨 颜色标记', level=2)
P('动词用橙色：', size=10.5)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
r1 = p.add_run('I '); r1.font.size = Pt(14); r1.font.name = 'Georgia'; r1.font.color.rgb = RGBColor(*BLUE); r1.bold = True
r2 = p.add_run('eat '); r2.font.size = Pt(14); r2.font.name = 'Georgia'; r2.font.color.rgb = RGBColor(*ORANGE); r2.bold = True
r3 = p.add_run('breakfast.'); r3.font.size = Pt(14); r3.font.name = 'Georgia'

doc.add_heading('📝 本课动词', level=2)
verbs = [
    ['eat', '吃', 'drink', '喝'],
    ['run', '跑', 'walk', '走'],
    ['play', '玩', 'read', '读'],
    ['sleep', '睡觉', 'sing', '唱歌'],
    ['go', '去', 'like', '喜欢'],
]
for row in verbs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.space_after = Pt(1)
    for j, w in enumerate(row):
        if j % 2 == 0:
            r = p.add_run(f'{w} '); r.font.size = Pt(11); r.font.name = 'Georgia'; r.font.bold = True; r.font.color.rgb = RGBColor(*ORANGE)
        else:
            r = p.add_run(f'({w})    '); r.font.size = Pt(10); r.font.color.rgb = RGBColor(*GRAY)

Divider()

doc.add_heading('✏️ 练习题', level=2)

ExTitle(1, '选择题')
ChooseCorrect([
    ('I ___ breakfast.', 'A. eat   B. eats', 'A'),
    ('We ___ soccer.', 'A. play   B. plays', 'A'),
    ('They ___ water.', 'A. drink   B. drinks', 'A'),
    ('You ___ books.', 'A. reads   B. read', 'B'),
])

ExTitle(2, '填空题')
FillBlank([
    ('I ___ (跑) fast.', ['run', 'runs'], 'run'),
    ('We ___ (玩) soccer.', ['play', 'plays'], 'play'),
    ('They ___ (读) books.', ['read', 'reads'], 'read'),
    ('You ___ (去) to school.', ['go', 'goes'], 'go'),
])

ExTitle(3, '组句练习')
BuildSentence([
    (['I', 'eat', 'breakfast'], 'I eat breakfast.'),
    (['play', 'We', 'soccer'], 'We play soccer.'),
    (['They', 'read', 'books'], 'They read books.'),
    (['drink', 'I', 'water'], 'I drink water.'),
])

ExTitle(4, '看图写句子')
WriteFromPic([
    ('🍳（早饭）', 'I / eat / breakfast'),
    ('⚽（足球）', 'We / play / soccer'),
    ('📚（读书）', ''),
    ('💧（喝水）', ''),
])

doc.add_page_break()

# ====================================================================
# 第6课：He eats / She goes（三单 -s）
# ====================================================================
doc.add_heading('第六课  He eats / She goes（他/她做动作）', level=1)

doc.add_heading('📖 知识讲解', level=2)

P('这一课非常重要！当主语是 He、She、It 的时候，动词要加 -s 或 -es。', size=11)

P('⚠️ 这是最容易出错的地方，我们慢慢学。', bold=True, size=11, color=(0xE6, 0x51, 0x00), space_before=6)

# 对比
P('先看对比：', bold=True, size=11, space_before=8)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.space_after = Pt(2)
r1 = p.add_run('I eat breakfast.     '); r1.font.size = Pt(12); r1.font.name = 'Georgia'
r2 = p.add_run('← eat 没有 s'); r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(*GRAY)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.space_after = Pt(2)
r1 = p.add_run('He eats breakfast.   '); r1.font.size = Pt(12); r1.font.name = 'Georgia'
r2 = p.add_run('← 有 s 了！'); r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)

P('')
P('规则：', bold=True, size=11)
P('看到 He / She / It 做主语 → 动词后面加 -s（或 -es）', size=11, indent=1.0)

Example('He eats breakfast.   eat → eats', size=11, indent=1.0)
Example('She goes to school.  go → goes', size=11, indent=1.0)
Example('It drinks water.    drink → drinks', size=11, indent=1.0)

Tip('-s 就像一个小尾巴，只有当主语是 He/She/It 的时候才出现。你可以把它想象成：He 和 She 喜欢在动词后面带一个小尾巴！')

doc.add_heading('📝 常见的 -s / -es 变化', level=2)

changes = [
    ['一般动词 +s', 'eat→eats', 'play→plays', 'read→reads', 'like→likes'],
    ['以 s, x, sh, ch, o 结尾 +es', 'go→goes', 'watch→watches', 'wash→washes', 'do→does'],
    ['以辅音+y 结尾 → ies', 'fly→flies', 'study→studies', '', ''],
]
add_table(['规则', '例1', '例2', '例3', '例4'], changes, [3.0, 2.5, 2.5, 2.5, 2.5])

Divider()

doc.add_heading('✏️ 练习题', level=2)

ExTitle(1, '选择题（选正确的动词）')
ChooseCorrect([
    ('He ___ breakfast.', 'A. eat   B. eats', 'B'),
    ('She ___ to school.', 'A. go   B. goes', 'B'),
    ('It ___ water.', 'A. drink   B. drinks', 'B'),
    ('He ___ soccer.', 'A. play   B. plays', 'B'),
    ('I ___ books.  (注意：I 不是 He!)', 'A. read   B. reads', 'A'),
    ('She ___ fast.', 'A. run   B. runs', 'B'),
])

ExTitle(2, '填空题（动词要不要加 s？）')
FillBlank([
    ('He ___ (吃) breakfast.', ['eat', 'eats'], 'eats'),
    ('I ___ (跑) fast.', ['run', 'runs'], 'run'),
    ('She ___ (去) to school.', ['go', 'goes'], 'goes'),
    ('They ___ (玩) soccer.', ['play', 'plays'], 'play'),
    ('It ___ (喝) water.', ['drink', 'drinks'], 'drinks'),
    ('We ___ (读) books.', ['read', 'reads'], 'read'),
])
Tip('小贴士：做题前先看主语！He/She/It → 加 s，I/You/We/They → 不加 s。')

ExTitle(3, '组句练习')
BuildSentence([
    (['He', 'eats', 'breakfast'], 'He eats breakfast.'),
    (['She', 'goes', 'to', 'school'], 'She goes to school.'),
    (['I', 'eat', 'breakfast'], 'I eat breakfast.'),
    (['plays', 'He', 'soccer'], 'He plays soccer.'),
])

ExTitle(4, '看图写句子')
WriteFromPic([
    ('👦🍳（男孩+早饭）', 'He / eats / breakfast'),
    ('👧🏫（女孩+学校）', 'She / goes / to / school'),
    ('👦⚽（男孩+足球）', ''),
    ('🧒🍳（我+早饭）', ''),
])

doc.add_page_break()

# ====================================================================
# 第7课：don't / doesn't（不...）
# ====================================================================
doc.add_heading('第七课  don\'t / doesn\'t（不...）', level=1)

doc.add_heading('📖 知识讲解', level=2)

P('想说不做什么事情，用 don\'t 或 doesn\'t。', size=11)

P('规则：', bold=True, size=11, space_before=6)
P('I / You / We / They → 用 don\'t', size=11, indent=1.0)
P('He / She / It → 用 doesn\'t', size=11, indent=1.0)

Example('I like fish.   →   I don\'t like fish.', size=11, indent=1.0)
P('  （我喜欢鱼）       （我不喜欢鱼）', size=9, color=GRAY, indent=1.0)

Example('He likes fish. →   He doesn\'t like fish.', size=11, indent=1.0)
P('  （他喜欢鱼）       （他不喜欢鱼）', size=9, color=GRAY, indent=1.0)

P('')
P('⚠️ 超级重要的规则！', bold=True, size=12, color=(0xE6, 0x51, 0x00), space_before=8)
P('doesn\'t 后面的动词不加 -s！', bold=True, size=12, indent=1.0, color=(0xE6, 0x51, 0x00))

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.space_before = Pt(4)
r1 = p.add_run('✓ He doesn\'t like fish.  '); r1.font.size = Pt(12); r1.font.name = 'Georgia'; r1.font.color.rgb = RGBColor(*GREEN)
r2 = p.add_run('（like 没有 s！）'); r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(*GREEN)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
r1 = p.add_run('✗ He doesn\'t likes fish. '); r1.font.size = Pt(12); r1.font.name = 'Georgia'; r1.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
r2 = p.add_run('（错！多了 s）'); r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)

Tip('一个简单的记法：doesn\'t 里面的 does 已经"带着 s"了，所以动词就不需要再加 s 了。就像一个人已经戴了帽子，不需要再戴一顶！')

Divider()

doc.add_heading('✏️ 练习题', level=2)

ExTitle(1, '选择题')
ChooseCorrect([
    ('I ___ like fish.', 'A. don\'t   B. doesn\'t', 'A'),
    ('He ___ like fish.', 'A. don\'t   B. doesn\'t', 'B'),
    ('She ___ play soccer.', 'A. don\'t   B. doesn\'t', 'B'),
    ('We ___ go to school.', 'A. don\'t   B. doesn\'t', 'A'),
])

ExTitle(2, '填空题（选 don\'t 或 doesn\'t）')
FillBlank([
    ('I ___ like fish.', ['don\'t', 'doesn\'t'], 'don\'t'),
    ('He ___ like dogs.', ['don\'t', 'doesn\'t'], 'doesn\'t'),
    ('She ___ eat breakfast.', ['don\'t', 'doesn\'t'], 'doesn\'t'),
    ('They ___ play soccer.', ['don\'t', 'doesn\'t'], 'don\'t'),
    ('It ___ drink water.', ['don\'t', 'doesn\'t'], 'doesn\'t'),
    ('You ___ run fast.', ['don\'t', 'doesn\'t'], 'don\'t'),
])

ExTitle(3, '改错题（把下面句子改成否定句）')
P('1. I like fish.  →  I _____ like fish.', size=11, indent=1.2, space_after=6)
P('2. He eats breakfast.  →  He _____ eat breakfast.', size=11, indent=1.2, space_after=6)
P('3. She plays soccer.  →  She _____ play soccer.', size=11, indent=1.2, space_after=6)
P('4. They go to school.  →  They _____ go to school.', size=11, indent=1.2, space_after=6)
Tip('注意第2题和第3题！eats 要变回 eat，plays 要变回 play！因为 doesn\'t 后面的动词不加 s！')

ExTitle(4, '组句练习')
BuildSentence([
    (['I', 'don\'t', 'like', 'fish'], 'I don\'t like fish.'),
    (['He', 'doesn\'t', 'like', 'fish'], 'He doesn\'t like fish.'),
    (['doesn\'t', 'She', 'play', 'soccer'], 'She doesn\'t play soccer.'),
])

doc.add_page_break()

# ====================================================================
# 第8课：can / can't（能/不能）
# ====================================================================
doc.add_heading('第八课  can / can\'t（能 / 不能）', level=1)

doc.add_heading('📖 知识讲解', level=2)

P('can 表示"能、会、可以"做某事。can\'t 表示"不能、不会"。', size=11)

P('最大的好处：can 后面的动词永远不变化！', bold=True, size=12, color=GREEN, space_before=6)

Example('I can swim.    （我会游泳。）', size=11, indent=1.0)
Example('He can swim.   （他会游泳。）', size=11, indent=1.0)
Example('She can swim.  （她会游泳。）', size=11, indent=1.0)
Example('They can swim. （他们会游泳。）', size=11, indent=1.0)

P('看到了吗？不管是 I、He、She、They，can 后面的 swim 都不变！', size=11, indent=1.0, space_before=4)

P('否定形式：', bold=True, size=11, space_before=8)
Example('I can\'t fly.   （我不能飞。）', size=11, indent=1.0)
Example('He can\'t swim. （他不会游泳。）', size=11, indent=1.0)

P('can\'t = can + not 的缩写。发音类似 "康特"。', size=10, color=GRAY, indent=1.0)

doc.add_heading('📝 本课动词', level=2)
can_verbs = ['swim 🏊 游泳', 'run 🏃 跑', 'jump 🤾 跳', 'sing 🎤 唱歌', 'dance 💃 跳舞', 'cook 🍳 做饭', 'fly 🕊️ 飞', 'read 📖 读']
for v in can_verbs:
    P(v, size=10.5, indent=1.0, space_after=1)

Divider()

doc.add_heading('✏️ 练习题', level=2)

ExTitle(1, '选择题')
ChooseCorrect([
    ('I ___ swim.', 'A. can   B. cans', 'A'),
    ('He ___ run fast.', 'A. can   B. cans', 'A'),
    ('Birds ___ fly.', 'A. can   B. are', 'A'),
    ('Dogs ___ fly.', 'A. can\'t   B. don\'t can', 'A'),
    ('She ___ sing.', 'A. can   B. is can', 'A'),
    ('He ___ drive.', 'A. can\'t   B. doesn\'t can', 'A'),
])
Tip('can 没有 cans 这种形式！不管主语是什么，can 永远不变。')

ExTitle(2, '填空题（选 can 或 can\'t）')
FillBlank([
    ('Fish ___ swim.', ['can', 'can\'t'], 'can'),
    ('Birds ___ fly.', ['can', 'can\'t'], 'can'),
    ('Dogs ___ fly.', ['can', 'can\'t'], 'can\'t'),
    ('A baby ___ drive.', ['can', 'can\'t'], 'can\'t'),
    ('I ___ read.', ['can', 'can\'t'], 'can'),
    ('A cat ___ sing.', ['can', 'can\'t'], 'can\'t'),
])

ExTitle(3, '组句练习')
BuildSentence([
    (['I', 'can', 'swim'], 'I can swim.'),
    (['He', 'can', 'run', 'fast'], 'He can run fast.'),
    (['can\'t', 'I', 'fly'], 'I can\'t fly.'),
    (['She', 'can', 'sing'], 'She can sing.'),
])

ExTitle(4, '回答问题（用 Yes, I can. 或 No, I can\'t.）')
P('1. Can you swim?  →  ____________________', size=11, indent=1.2, space_after=6)
P('2. Can you fly?   →  ____________________', size=11, indent=1.2, space_after=6)
P('3. Can you read?  →  ____________________', size=11, indent=1.2, space_after=6)
P('4. Can you sing?  →  ____________________', size=11, indent=1.2, space_after=6)

doc.add_page_break()

# ====================================================================
# 第9课：正在进行时（be + ing）
# ====================================================================
doc.add_heading('第九课  正在做...（be + -ing）', level=1)

doc.add_heading('📖 知识讲解', level=2)

P('想表达"正在做什么"，用 am/is/are + 动词-ing。', size=11)

P('公式：主语 + am/is/are + 动词-ing', bold=True, size=12, space_before=6)

Example('I am reading.     （我正在读书。）', size=11, indent=1.0)
Example('He is sleeping.   （他正在睡觉。）', size=11, indent=1.0)
Example('She is running.   （她正在跑步。）', size=11, indent=1.0)
Example('They are playing. （他们正在玩。）', size=11, indent=1.0)

P('')
P('am/is/are 的选择规则和第一课一样：', size=10.5, indent=1.0)
P('I → am    He/She → is    You/We/They → are', size=10.5, indent=2.0)

doc.add_heading('📝 -ing 的变化规则', level=2)

ing_rules = [
    ['一般动词 +ing', 'eat→eating', 'play→playing', 'read→reading'],
    ['以 e 结尾 → 去 e +ing', 'make→making', 'write→writing', 'dance→dancing'],
    ['短元音+辅音 → 双写+ing', 'run→running', 'swim→swimming', 'sit→sitting'],
]
add_table(['规则', '例1', '例2', '例3'], ing_rules, [3.5, 3.5, 3.5, 3.5])

Divider()

doc.add_heading('✏️ 练习题', level=2)

ExTitle(1, '选择题')
ChooseCorrect([
    ('I ___ reading.', 'A. am   B. is   C. are', 'A'),
    ('He ___ sleeping.', 'A. am   B. is   C. are', 'B'),
    ('They ___ playing.', 'A. am   B. is   C. are', 'C'),
    ('She is ___.', 'A. run   B. running', 'B'),
    ('I am ___ water.', 'A. drink   B. drinking', 'B'),
])

ExTitle(2, '填空题')
FillBlank([
    ('I am ___ (读).', ['read', 'reading'], 'reading'),
    ('He is ___ (睡).', ['sleep', 'sleeping'], 'sleeping'),
    ('She is ___ (跑).', ['run', 'running'], 'running'),
    ('They are ___ (玩).', ['play', 'playing'], 'playing'),
    ('We are ___ (唱歌).', ['sing', 'singing'], 'singing'),
    ('It is ___ (喝) water.', ['drink', 'drinking'], 'drinking'),
])

ExTitle(3, '组句练习')
BuildSentence([
    (['I', 'am', 'reading'], 'I am reading.'),
    (['He', 'is', 'sleeping'], 'He is sleeping.'),
    (['are', 'They', 'playing'], 'They are playing.'),
    (['She', 'is', 'running'], 'She is running.'),
])

ExTitle(4, '看图写句子')
WriteFromPic([
    ('📖（读书）I am...', ''),
    ('😴（睡觉）He is...', ''),
    ('⚽（踢球）They are...', ''),
    ('🏃（跑步）She is...', ''),
])

doc.add_page_break()

# ====================================================================
# 第10课：介词（方位词 in / on / under）
# ====================================================================
doc.add_heading('第十课  方位词（in / on / under...）', level=1)

doc.add_heading('📖 知识讲解', level=2)

P('方位词告诉我们东西在哪里。', size=11)

P('最常用的三个方位词：', size=11, space_before=6)

# 三个介词配说明
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.space_before = Pt(6)
r1 = p.add_run('in  '); r1.font.size = Pt(14); r1.font.name = 'Georgia'; r1.font.color.rgb = RGBColor(*PURPLE); r1.bold = True
r2 = p.add_run('= 在里面'); r2.font.size = Pt(11)

Example('The cat is in the box.   猫在盒子里面。', size=11, indent=1.0)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.space_before = Pt(6)
r1 = p.add_run('on  '); r1.font.size = Pt(14); r1.font.name = 'Georgia'; r1.font.color.rgb = RGBColor(*PURPLE); r1.bold = True
r2 = p.add_run('= 在上面'); r2.font.size = Pt(11)

Example('The cat is on the box.    猫在盒子上面。', size=11, indent=1.0)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.space_before = Pt(6)
r1 = p.add_run('under  '); r1.font.size = Pt(14); r1.font.name = 'Georgia'; r1.font.color.rgb = RGBColor(*PURPLE); r1.bold = True
r2 = p.add_run('= 在下面'); r2.font.size = Pt(11)

Example('The cat is under the box. 猫在盒子下面。', size=11, indent=1.0)

doc.add_heading('🎨 颜色标记', level=2)
P('方位词用紫色（小帮手词）：', size=10.5)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.space_before = Pt(4)
r1 = p.add_run('The cat is '); r1.font.size = Pt(13); r1.font.name = 'Georgia'
r2 = p.add_run('in'); r2.font.size = Pt(13); r2.font.name = 'Georgia'; r2.font.color.rgb = RGBColor(*PURPLE); r2.bold = True
r3 = p.add_run(' the box.'); r3.font.size = Pt(13); r3.font.name = 'Georgia'

Divider()

doc.add_heading('✏️ 练习题', level=2)

ExTitle(1, '选择题')
ChooseCorrect([
    ('🐱📦（猫在盒子里）The cat is ___ the box.', 'A. in   B. on   C. under', 'A'),
    ('🐱📦（猫在盒子上）The cat is ___ the box.', 'A. in   B. on   C. under', 'B'),
    ('🐱📦（猫在盒子下）The cat is ___ the box.', 'A. in   B. on   C. under', 'C'),
    ('📚🪑（书在椅子上）The book is ___ the chair.', 'A. in   B. on   C. under', 'B'),
    ('🐕🪑（狗在椅子下）The dog is ___ the chair.', 'A. in   B. on   C. under', 'C'),
])

ExTitle(2, '填空题（选 in / on / under）')
FillBlank([
    ('The cat is ___ the box. (里面)', ['in', 'on', 'under'], 'in'),
    ('The cat is ___ the box. (上面)', ['in', 'on', 'under'], 'on'),
    ('The cat is ___ the box. (下面)', ['in', 'on', 'under'], 'under'),
    ('The book is ___ the desk. (上面)', ['in', 'on', 'under'], 'on'),
    ('The dog is ___ the chair. (下面)', ['in', 'on', 'under'], 'under'),
    ('The pen is ___ the bag. (里面)', ['in', 'on', 'under'], 'in'),
])

ExTitle(3, '组句练习')
BuildSentence([
    (['The', 'cat', 'is', 'in', 'the', 'box'], 'The cat is in the box.'),
    (['The', 'cat', 'is', 'on', 'the', 'box'], 'The cat is on the box.'),
    (['The', 'cat', 'is', 'under', 'the', 'box'], 'The cat is under the box.'),
    (['The', 'book', 'is', 'on', 'the', 'desk'], 'The book is on the desk.'),
])

ExTitle(4, '看图写句子')
WriteFromPic([
    ('🐱📦（猫在盒子里面）', 'in'),
    ('🐱📦（猫在盒子上面）', 'on'),
    ('🐕🪑（狗在椅子下面）', ''),
    ('📚🪑（书在桌子上面）', ''),
])

doc.add_page_break()

# ====================================================================
# 第11课：have / has（有）
# ====================================================================
doc.add_heading('第十一课  have / has（有）', level=1)

doc.add_heading('📖 知识讲解', level=2)

P('have 和 has 都表示"有"。选哪个取决于主语是谁。', size=11)

P('规则和 am / is / are 很像！', bold=True, size=11, space_before=6)

# 对照表
P('')
add_table(['主语', '用哪个', '例句'], [
    ['I', 'have', 'I have a dog.'],
    ['You', 'have', 'You have a book.'],
    ['We', 'have', 'We have a ball.'],
    ['They', 'have', 'They have a cat.'],
    ['He', 'has', 'He has a dog.'],
    ['She', 'has', 'She has a cat.'],
    ['It', 'has', 'It has a tail.'],
], [2.5, 2.0, 5.5])

P('')
P('发现规律了吗？', bold=True, size=11)
P('I / You / We / They → have（不变）', size=11, indent=1.0)
P('He / She / It → has（加 s）', size=11, indent=1.0)
Tip('这和 am/is/are 的规则基本一样！只是把 are 换成了 have，把 is 换成了 has。记住：he/she/it 后面用 has（和 is 一样带 s！）')

Divider()

doc.add_heading('✏️ 练习题', level=2)

ExTitle(1, '选择题')
ChooseCorrect([
    ('I ___ a dog.', 'A. have   B. has', 'A'),
    ('He ___ a cat.', 'A. have   B. has', 'B'),
    ('She ___ a book.', 'A. have   B. has', 'B'),
    ('They ___ a ball.', 'A. have   B. has', 'A'),
    ('We ___ a friend.', 'A. have   B. has', 'A'),
    ('It ___ a tail.', 'A. have   B. has', 'B'),
])

ExTitle(2, '填空题（填 have 或 has）')
FillBlank([
    ('I ___ a pen.', ['have', 'has'], 'have'),
    ('He ___ a bag.', ['have', 'has'], 'has'),
    ('She ___ blue eyes.', ['have', 'has'], 'has'),
    ('We ___ a dog.', ['have', 'has'], 'have'),
    ('They ___ books.', ['have', 'has'], 'have'),
    ('It ___ a long tail.', ['have', 'has'], 'has'),
    ('You ___ a friend.', ['have', 'has'], 'have'),
    ('My cat ___ green eyes.', ['have', 'has'], 'has'),
])

ExTitle(3, '组句练习')
BuildSentence([
    (['I', 'have', 'a', 'dog'], 'I have a dog.'),
    (['He', 'has', 'a', 'cat'], 'He has a cat.'),
    (['She', 'has', 'a', 'book'], 'She has a book.'),
    (['They', 'have', 'a', 'ball'], 'They have a ball.'),
])

ExTitle(4, '看图写句子')
WriteFromPic([
    ('🐕（我有一只狗）I...', ''),
    ('🐱（她有一只猫）She...', ''),
    ('📚（他有一本书）He...', ''),
    ('⚽（我们有一个球）We...', ''),
])

doc.add_page_break()

# ====================================================================
# 第12课：综合复习
# ====================================================================
doc.add_heading('第十二课  综合复习', level=1)

P('恭喜你学完了所有 8 个单元！最后一起来做综合练习。', size=11, space_after=8)

doc.add_heading('📋 语法速查表', level=2)

review_table = [
    ['语法点', '规则', '例句'],
    ['Be动词', 'I am / He-She is / You-We-They are', 'I am happy. He is tall.'],
    ['a / an', '辅音前 a / 元音前 an', 'a cat / an apple'],
    ['一般现在时', 'He-She-It + V+s / 其他不变', 'He eats. / I eat.'],
    ['否定', 'don\'t / doesn\'t + V原形', 'He doesn\'t like fish.'],
    ['can', 'can + V原形（永远不变）', 'I can swim. / He can swim.'],
    ['进行时', 'am/is/are + V-ing', 'I am reading.'],
    ['方位词', 'in / on / under', 'The cat is in the box.'],
    ['have/has', 'He-She-It has / 其他 have', 'He has a dog. / I have a cat.'],
]
add_table(['语法点', '规则', '例句'], review_table, [2.5, 5.0, 5.5])

Divider()

doc.add_heading('✏️ 综合练习题', level=2)

ExTitle(1, '混合选择题')
ChooseCorrect([
    ('I ___ a student.', 'A. am   B. is   C. are', 'A'),
    ('He ___ a dog.', 'A. have   B. has', 'B'),
    ('They ___ playing soccer.', 'A. am   B. is   C. are', 'C'),
    ('She ___ sing.', 'A. can   B. cans', 'A'),
    ('The cat is ___ the box. (里面)', 'A. in   B. on   C. under', 'A'),
    ('He ___ breakfast every day.', 'A. eat   B. eats', 'B'),
    ('She ___ like fish.', 'A. don\'t   B. doesn\'t', 'B'),
    ('___ you swim?', 'A. Can   B. Do   C. Are', 'A'),
])

ExTitle(2, '改错题（找出错误并在横线上写出正确句子）')
P('1. He have a dog.     →  ________________________________', size=11, indent=1.2, space_after=4)
P('2. She can sings.     →  ________________________________', size=11, indent=1.2, space_after=4)
P('3. I is happy.        →  ________________________________', size=11, indent=1.2, space_after=4)
P('4. He don\'t like fish. →  ________________________________', size=11, indent=1.2, space_after=4)
P('5. They is playing.   →  ________________________________', size=11, indent=1.2, space_after=4)

ExTitle(3, '组句练习')
BuildSentence([
    (['He', 'is', 'reading', 'a', 'book'], 'He is reading a book.'),
    (['The', 'cat', 'is', 'under', 'the', 'table'], 'The cat is under the table.'),
    (['She', 'has', 'a', 'cat'], 'She has a cat.'),
    (['They', 'can', 'swim'], 'They can swim.'),
])

ExTitle(4, '翻译练习（把中文意思写成英语句子）')
P('1. 我是一个学生。        →  ________________________________', size=11, indent=1.2, space_after=4)
P('2. 他有一只猫。          →  ________________________________', size=11, indent=1.2, space_after=4)
P('3. 她正在跑步。          →  ________________________________', size=11, indent=1.2, space_after=4)
P('4. 猫在盒子里面。        →  ________________________________', size=11, indent=1.2, space_after=4)
P('5. 我不会飞。            →  ________________________________', size=11, indent=1.2, space_after=4)

doc.add_page_break()

# ====================================================================
# 尾页
# ====================================================================
doc.add_heading('🎉 祝贺你！', level=1)

P('你完成了全部 12 课的学习！你学了：', size=12, space_after=8)

achievements = [
    '✅ 用 am / is / are 说"是谁"、"怎么样"',
    '✅ 用 a / an 说"一个什么东西"',
    '✅ 用动词说"做什么事情"',
    '✅ 用 don\'t / doesn\'t 说"不做什么"',
    '✅ 用 can / can\'t 说"能做什么"',
    '✅ 用 am/is/are + ing 说"正在做什么"',
    '✅ 用 in / on / under 说"在哪里"',
    '✅ 用 have / has 说"有什么"',
]
for a in achievements:
    P(a, size=11, indent=1.5, space_after=3)

P('')
P('你已经掌握了英语最基础、最重要的语法！继续加油！🌟', size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)

# ============ 保存 ============
output_path = r'C:\西安备课\语法训练\ASD语法训练-学生用书.docx'
doc.save(output_path)
print(f'Done: {output_path}')
desktop = r'C:\Users\25152\Desktop\ASD语法训练-学生用书.docx'
shutil.copy2(output_path, desktop)
print(f'Copied to desktop')
