# -*- coding: utf-8 -*-
"""生成 ASD 青少年英语语法训练 Word 教案"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ============ 页面设置 ============
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============ 样式定义 ============
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(6)

# 标题样式
for i in range(1, 4):
    h_style = doc.styles[f'Heading {i}']
    h_font = h_style.font
    h_font.name = '微软雅黑'
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if i == 1:
        h_font.size = Pt(22)
        h_font.color.rgb = RGBColor(0x4A, 0x90, 0xD9)
        h_font.bold = True
    elif i == 2:
        h_font.size = Pt(16)
        h_font.color.rgb = RGBColor(0x5C, 0xB8, 0x5C)
        h_font.bold = True
    elif i == 3:
        h_font.size = Pt(13)
        h_font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        h_font.bold = True

# ============ 辅助函数 ============
def add_para(text, bold=False, size=None, color=None, align=None, space_after=None, indent=None):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p

def add_tip(text):
    """添加教学提示（缩进+灰色框）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('💡 ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_warning(text):
    """添加注意事项"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('⚠️ ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_table(headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表后空行
    return table

def add_step_box(step_num, step_title, content, bg_color=None):
    """添加教学步骤框"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'Step {step_num}: {step_title}')
    run.bold = True
    run.font.size = Pt(12)
    if step_num == 1:
        run.font.color.rgb = RGBColor(0x4A, 0x90, 0xD9)
    elif step_num == 2:
        run.font.color.rgb = RGBColor(0x5C, 0xB8, 0x5C)
    elif step_num == 3:
        run.font.color.rgb = RGBColor(0xF0, 0xAD, 0x4E)
    elif step_num == 4:
        run.font.color.rgb = RGBColor(0x9B, 0x59, 0xB6)
    elif step_num == 5:
        run.font.color.rgb = RGBColor(0xE9, 0x1E, 0x63)

    if content:
        add_para(content, size=10, indent=0.8, color=(0x55,0x55,0x55))

def add_exercise(label, items, indent=0.8):
    """添加练习题"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label + '  ')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x9B, 0x59, 0xB6)
    for i, item in enumerate(items[:8]):  # 最多8个
        run2 = p.add_run(f'{item}    ')
        run2.font.size = Pt(10)

# ====================================================================
# 封面
# ====================================================================
for _ in range(4):
    doc.add_paragraph()

add_para('🧩', size=36, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('ASD 青少年英语语法训练', size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x4A, 0x90, 0xD9))
add_para('循序渐进 · 视觉引导 · 情绪安全', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))
doc.add_paragraph()
add_para('8 级 24 关 · 完整教案', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x66, 0x66, 0x66))
add_para(f'版本 1.0 · {datetime.date.today().strftime("%Y年%m月%d日")}', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))

doc.add_page_break()

# ====================================================================
# 目录页
# ====================================================================
doc.add_heading('目录', level=1)
toc_items = [
    '一、ASD 学习者与语法习得',
    '二、设计核心原则',
    '三、语法等级体系总览（8 级 24 关）',
    '四、五步教学法详解',
    '五、语法色彩编码系统',
    '六、渐隐提示引擎',
    '七、Level 1 详细教案 — Be 动词基础',
    '八、Level 2 详细教案 — 冠词 + 名词复数',
    '九、Level 3 详细教案 — 一般现在时（无三单）',
    '十、Level 4 详细教案 — 一般现在时（三单）',
    '十一、Level 5 教案 — 情态动词 Can',
    '十二、Level 6 教案 — 现在进行时',
    '十三、Level 7 教案 — 介词方位',
    '十四、Level 8 教案 — Have / Has',
    '十五、情绪安全反馈规范',
    '十六、教师使用指南',
]
for item in toc_items:
    add_para(item, size=11, space_after=4)

doc.add_page_break()

# ====================================================================
# 一、ASD 学习者与语法习得
# ====================================================================
doc.add_heading('一、ASD 学习者与语法习得', level=1)

doc.add_heading('1.1 核心挑战', level=2)

challenges = [
    ['执行功能弱', '难以规划、切换任务', '不能同时处理多个语法规则；需一次只学一个点'],
    ['工作记忆有限', '信息保持时间短', '句子长了就忘记开头；需要超短句起步（3-5词）'],
    ['抽象思维困难', '偏好具体、可视化信息', '"时态""语序"等概念需视觉锚点，不能纯文字解释'],
    ['泛化能力弱', '学会A不等于会B', '每种句型模式必须显式教过才能迁移'],
    ['感官敏感', '对颜色、声音、动画敏感', '界面极简、安静、可预测；禁止闪烁和突然音效'],
    ['焦虑高发', '害怕犯错', '一次失败可能拒绝继续；必须零负面反馈'],
    ['刻板行为', '可能反复做同一题', '系统需内置温和推进机制'],
    ['语言处理差异', '功能词（the/a/is）常被忽略', '功能词需要高亮标记才能被注意到'],
]
add_table(['挑战维度', '具体表现', '对语法学习的影响'], challenges, [3.0, 4.5, 7.5])

doc.add_heading('1.2 ASD 学习者的优势（可利用）', level=2)

strengths = [
    ['模式识别能力强', '一旦建立"规则→例句→视觉标记"的映射，识别速度可能超常'],
    ['喜欢重复和可预测性', '恰好匹配语法练习需要的大量重复'],
    ['对规则忠诚', '学会规则后会严格应用，不容易随性犯错'],
    ['视觉学习偏好', '色彩编码、图式、表格是天然的教学工具'],
    ['专注力（感兴趣时）', '对特定主题的持续注意力可能远超常人'],
]
add_table(['优势', '教学利用方式'], strengths, [4.5, 10.5])

add_warning('ASD 不是一种缺陷，而是一种不同的认知方式。教学目标是搭建桥梁，而非"修复"。')

doc.add_page_break()

# ====================================================================
# 二、设计核心原则
# ====================================================================
doc.add_heading('二、设计核心原则', level=1)

principles = [
    ('原则一：一次一个语法点（One Rule at a Time）',
     '不把 am/is/are 和所有人称代词一起教。先只练 "I am..."，熟练后再加 "He is..."，最后才混合。',
     '❌ 错误：同时教 am/is/are + 人称代词 + 形容词\n✅ 正确：第1关只练 I am → 第2关 He/She is → 第3关 You/They are → 第4关混合'),

    ('原则二：视觉语法标记（Visual Grammar Markers）',
     '每种语法成分有固定颜色，整个系统保持一致。ASD 学生通过颜色识别语法功能，而非通过抽象术语。',
     '蓝色 = 主语 · 绿色 = Be动词 · 橙色 = 动作动词 · 紫色 = 功能词 · 黄色 = 词尾变化'),

    ('原则三：具体先于抽象（Concrete Before Abstract）',
     '每个语法点都从图片+例句开始，再逐渐去掉图片、去掉颜色、去掉提示，最终独立产出。',
     '图片+例句 → 色块标记 → 填空选择 → 拖拽组句 → 独立写句'),

    ('原则四：零负面反馈（Emotional Safety）',
     '绝对不用"错了""不对""❌""红色叉号"。错误时只自动显示正确答案（黄色闪烁1.5秒），不评价、不对比、不排名。',
     '正确：⭐ Good job! · 不正确：自动显示答案，无评价 · 没有分数，只有星星收集'),

    ('原则五：超量重复 + 间隔复习',
     '新语法点当天做 8-12 次同一模式练习（massed practice），正确率 ≥85% 后进入间隔复习（1天后→3天后→7天后）。',
     '新学：8-12次同模式 → 间隔1天复习4题 → 间隔3天复习3题 → 间隔7天复习2题 → 标记"已掌握"'),

    ('原则六：可预测的流程（Predictable Routine）',
     '每关结构完全相同，只是语法内容不同。ASD 学生一旦熟悉流程，认知负荷大幅降低。',
     '每关固定：SEE → COLOR → MATCH → BUILD → USE，16-20题，8-15分钟'),
]

for title, desc, example in principles:
    doc.add_heading(title, level=3)
    add_para(desc, size=10.5)
    add_tip(example)

doc.add_page_break()

# ====================================================================
# 三、语法等级体系总览
# ====================================================================
doc.add_heading('三、语法等级体系总览（8 级 24 关）', level=1)

add_para('以下是完整的 8 级 24 关语法路线图。每个 Level 包含 3 关，每关 16-20 题，约需 8-15 分钟完成。', size=10.5)

levels_overview = [
    ['Level 1 🟢', 'Be 动词基础', '1-1 I am ...\n1-2 He/She is ...\n1-3 You/We/They are ...', '英语句子的"骨架"，一切的基础'],
    ['Level 2 🟣', '冠词 + 名词复数', '2-1 a vs an\n2-2 a/an vs the\n2-3 名词复数 -s/-es/-ies', '名词的基本包装，高度可视化'],
    ['Level 3 🟠', '一般现在时\n（无三单）', '3-1 I/You/We/They + V原形\n3-2 否定 don\'t\n3-3 疑问 Do...?', '动作表达的第一步，先不碰三单难点'],
    ['Level 4 🟡', '一般现在时\n（三单）', '4-1 He/She/It + V+s/es\n4-2 否定 doesn\'t\n4-3 疑问 Does...?', '最难点！拆成3关，用"偷走-s"动画'],
    ['Level 5 🟣', '情态动词 Can', '5-1 I can / can\'t\n5-2 Can you...? 问答\n5-3 He/She can', '动词永远不变，给学生信心'],
    ['Level 6 🟠', '现在进行时', '6-1 I am + V-ing\n6-2 He/She is + V-ing\n6-3 They are + V-ing', '用已会的be动词+已会的动词+-ing'],
    ['Level 7 🟣', '介词方位', '7-1 in / on / under\n7-2 next to / behind\n7-3 There is/are + 介词', 'ASD优势领域——空间思维，最可视化'],
    ['Level 8 🔵', 'Have / Has', '8-1 I/You/We/They have\n8-2 He/She/It has\n8-3 综合复习', '收尾，类比已学的am/is/are规则'],
]
add_table(['级别', '内容', '关卡', '设计理由'], levels_overview, [2.0, 3.0, 4.5, 5.5])

add_tip('为什么选这8个语法点？这8个点覆盖日常英语 70%+ 的句子结构；每个点都有对应图片/场景，不依赖抽象解释；有严格的先修依赖关系（L1→L2→L3→L4）；24关的进度条让 ASD 学生感到"可以完成"。')

doc.add_page_break()

# ====================================================================
# 四、五步教学法详解
# ====================================================================
doc.add_heading('四、五步教学法详解', level=1)

add_para('每关内部采用固定的五步教学法。ASD 学生一旦适应流程，就不再需要理解"新指令"。五种题型在每关中按固定顺序出现。', size=10.5)

steps_detail = [
    ('Step 1: SEE 👀 — 看例句',
     '展示一张清晰的图片 + 一个正确的英文句子。系统慢速朗读（0.4x），学生被动观察 10-15 秒。',
     '目的：建立"画面 → 英文句子"的直接映射，不解释语法规则。',
     '示例：🧒男孩图片 → "I am a boy." 🔊'),
    ('Step 2: COLOR 🎨 — 色块标注',
     '与 Step 1 相同的图片+句子，但加上语法颜色标记。颜色块逐个缓慢出现（每个 0.8 秒），引导学生注意功能词。',
     '目的：让 ASD 学生"看到"语法结构。功能词（is/the/a）用显眼颜色标记，因为这些词常被 ASD 学生跳过。',
     '示例：🔵I（蓝色）→ 🟢am（绿色）→ 🟣a（紫色）→ 🔵boy（蓝色），逐个出现'),
    ('Step 3: MATCH 🔗 — 选择填空',
     '2选1或3选1的选择题。提示等级从全开（选项有颜色标记）到渐隐（选项无颜色）。',
     '目的：在高度结构化的选择中建立信心。错误自动显示答案，不评价。',
     '示例：I ___ a student.  [ am / is ]'),
    ('Step 4: BUILD 🧱 — 拖拽组句',
     '屏幕底部有乱序词块，拖到上方横线组成句子。词块颜色从有色到无色渐隐。',
     '目的：理解英语语序（主-谓-宾），通过动手操作加深记忆。',
     '示例：底部词块 [am] [I] [happy] → 拖到上方组成 I am happy.'),
    ('Step 5: USE ✏️ — 看图写句',
     '只给图片，不给任何文字提示。学生独立写出完整句子。宽松评分（大小写、句号不扣分）。',
     '目的：最终目标——独立产出。这是检验是否真正掌握的关卡。',
     '示例：🧒 → 学生写出 "I am a boy."'),
]

for title, desc, purpose, example in steps_detail:
    doc.add_heading(title, level=3)
    add_para(desc, size=10.5)
    add_para(purpose, size=10, indent=0.8, color=(0x5C, 0xB8, 0x5C))
    add_tip(example)

doc.add_page_break()

# ====================================================================
# 五、语法色彩编码系统
# ====================================================================
doc.add_heading('五、语法色彩编码系统', level=1)

add_para('这是整个系统的核心视觉工具。每种语法成分有固定的颜色标记，在 Level 1-8 中保持完全一致。ASD 学生通过颜色来"看到"语法结构，而非通过术语理解。', size=10.5)

colors = [
    ['🔵 蓝色', '#4A90D9', '主语（名词/代词）', 'I, you, he, she, it, we, they, cat, dog, boy'],
    ['🟢 绿色', '#5CB85C', 'Be 动词 / Have', 'am, is, are, have, has'],
    ['🟠 橙色', '#F0AD4E', '动作动词', 'eat, drink, run, play, read, sleep, swim, sing'],
    ['🟣 紫色', '#9B59B6', '功能词（小帮手词）', 'the, a, an, can, in, on, under, do, don\'t, doesn\'t'],
    ['🟡 黄色加粗', '#FFC107', '词尾屈折变化', '-s, -es, -ies, -ing, -ed'],
    ['🔴 粉色', '#E91E63', '时间标记', 'now, every day, right now'],
]
add_table(['颜色', '色值', '语法成分', '示例词'], colors, [2.5, 2.0, 3.5, 7.0])

add_tip('使用规则：① 颜色编码在整个学习过程中保持一致，不随关卡变化。② Step 2（COLOR）中颜色逐个出现，让学生有时间"消化"每个词。③ 随着渐隐等级提升，颜色逐渐褪去，最终学生看到的是普通黑白文字。')

doc.add_page_break()

# ====================================================================
# 六、渐隐提示引擎
# ====================================================================
doc.add_heading('六、渐隐提示引擎', level=1)

add_para('渐隐（Fading）是 ASD 教学中最有效的策略之一。从最大支持到完全独立，让学生在不知不觉中建立能力。', size=10.5)

fading = [
    ['L4', '全提示', '选项保留颜色标记 + 语法对照表可见 + 💡按钮亮', '新语法点前 3 题'],
    ['L3', '半提示', '选项颜色去掉 + 对照表仍可见 + 💡按钮亮', '连续 2 题正确后升级'],
    ['L2', '弱提示', '对照表折叠隐藏 + 💡按钮暗（仍可点击）', '再连续 2 题正确后升级'],
    ['L1', '无提示', '无颜色、无对照表、无💡按钮', '再连续 2 题正确后升级'],
    ['L0', '独立', '完全独立产出（仅用于 Step 5 USE）', '本关最终目标'],
]
add_table(['等级', '名称', '效果', '触发条件'], fading, [1.2, 2.0, 7.0, 4.8])

doc.add_heading('升级 / 降级规则', level=3)
add_para('• 升级：连续 2 题正确 → 提示等级 -1（L4→L3→L2→L1→L0）', size=10.5)
add_para('• 降级：连续 2 题错误 → 提示等级 +1（L1→L2→L3→L4）', size=10.5)
add_para('• 题型切换时重置：从 MATCH 换到 BUILD 时，提示回到 L3（因为 BUILD 比 MATCH 难）', size=10.5)
add_para('• L0 只适用于 Step 5 USE：这是终极目标——没有任何帮助，看图独立写句', size=10.5)

doc.add_page_break()

# ====================================================================
# 七~十四：各级教案
# ====================================================================

# ---- Level 1 ----
doc.add_heading('七、Level 1 详细教案 — Be 动词基础', level=1)

add_para('Level 1 是整个语法体系的地基。学生需要在这里建立"主语 + Be动词 + 表语"的基本句型意识。分 3 关，每关聚焦一个人称组，最后一关混合。', size=10.5)

# 1-1
doc.add_heading('Level 1-1: I am ...（我是...）', level=2)
add_para('核心句型：I am + 身份/感受/描述', bold=True, size=10.5)
add_para('词汇准备：a boy, a girl, a student, a teacher, happy, sad, hungry, tired, tall, short, big, small', size=10)

add_step_box(1, 'SEE — 看例句（2题）', None)
add_exercise('题1-2:', ['🧒→I am a boy.', '😊→I am happy.'])

add_step_box(2, 'COLOR — 色块标注（2题）', None)
add_exercise('题3-4:', ['🧒→🔵I 🟢am 🟣a 🔵boy.', '😊→🔵I 🟢am 🔵happy.'])
add_tip('颜色逐个出现，每个间隔 0.8 秒。教师可旁白："看到绿色的 am 了吗？I 后面总是跟着 am。"')

add_step_box(3, 'MATCH — 选择填空（4题）', None)
add_exercise('题5-8:', ['🧒 I ___ a boy. [am✓/is]', '😊 I ___ happy. [am✓/are]', '😢 I ___ sad. [is/am✓]', '🧑 I ___ a student. [are/am✓]'])
add_tip('前2题选项保留颜色（L4），后2题去掉颜色（L3）。')

add_step_box(4, 'BUILD — 拖拽组句（4题）', None)
add_exercise('题9-12:', ['😊→[am][I][happy]→I am happy.', '🧒→[a][am][I][boy]→I am a boy.', '📏→[tall][I][am]→I am tall.', '🍽→[hungry][am][I]→I am hungry.'])
add_tip('前2题词块有颜色（L4），后2题无颜色（L3）。')

add_step_box(5, 'USE — 看图写句（4题）', None)
add_exercise('题13-16:', ['🧒→(写)I am a boy.', '😊→I am happy.', '😢→I am sad.', '📏→I am tall.'])
add_tip('题13-14有💡提示按钮（L2），题15-16无提示（L1）。')

add_warning('如果学生在题目中选了 "I is..." 或 "I are..."，不要纠正。等待 1.5 秒后自动显示正确答案 "I am..."，然后下一题。重复接触 8 次后，正确率通常达到 85%+。')

# 1-2
doc.add_heading('Level 1-2: He is / She is ...（他是/她是...）', level=2)
add_para('核心句型：He/She + is + 表语', bold=True, size=10.5)
add_para('关键教学点：am → is 的动词变化。引入"对比卡"环节。', size=10.5)

add_step_box(1, 'SEE — 看例句（2题）', None)
add_exercise('题1-2:', ['👦→He is a boy.', '👧→She is a girl.'])

add_step_box('1.5', '🔄 对比观察（2题，Level 1-2 特有环节！）', None)
add_para('同时展示两张卡片，让学生观察 am→is 的变化：', size=10)
add_exercise('题3-4:', ['🧒 I am happy. → 👦 He is happy. (am变is!)', '🧒 I am tall. → 👧 She is tall. (am变is!)'])
add_tip('关键设计：两张卡片并列展示，am 和 is 都用绿色闪烁 2 秒，让 ASD 学生"看到"变化。可以说："I 用 am，He 用 is。看出来了吗？"')

add_step_box(2, 'COLOR（2题）', None)
add_exercise('题5-6:', ['👦→🔵He 🟢is 🔵tall.', '👧→🔵She 🟢is 🔵happy.'])

add_step_box(3, 'MATCH（4题，选 am/is）', None)
add_exercise('题7-10:', ['👦 He ___ a boy. [am/is✓]', '👧 She ___ a girl. [am/is✓]', '👦 He ___ tall. [is✓/am]', '👧 She ___ happy. [am→自动显示is]'])
add_tip('题10 是故意设计——题目提供 am 和 is 两个选项，答案却是 is。如果选 am，1.5 秒后自动显示 is，不批评。这叫"发现学习"。')

add_step_box(4, 'BUILD（4题）', None)
add_exercise('题11-14:', ['👦→[is][He][tall]', '👧→[happy][She][is]', '👦→[a][is][He][boy]', '👧→[girl][She][is][a]'])

add_step_box(5, 'USE（4题）', None)
add_exercise('题15-18:', ['👦→He is a boy.', '👧→She is a girl.', '👦😊→He is happy.', '👧📏→She is tall.'])

# 1-3
doc.add_heading('Level 1-3: You are / We are / They are', level=2)
add_para('核心句型：You/We/They + are + 表语', bold=True, size=10.5)
add_para('关键教学点：引入完整的人称-动词对照表，作为"认知支架"。本关所有题目中对照表始终可见（L4-L3 阶段）。', size=10.5)

add_para('📋 对照表（本关始终展示）：', bold=True, size=10.5)
add_table(['主语', 'Be动词', '状态'], [
    ['I', 'am', '已学 ✓'],
    ['He / She', 'is', '已学 ✓'],
    ['You', 'are', '🆕 新学！'],
    ['We', 'are', '🆕 新学！'],
    ['They', 'are', '🆕 新学！'],
], [4.0, 4.0, 4.0])

add_step_box(3, 'MATCH（6题，3选1 am/is/are）', None)
add_exercise('题7-12:', ['👥 You ___ students. [am/is/are✓]', '👥 We ___ happy. [am/is/are✓]', '👥 They ___ tall. [am/is/are✓]',
    '🧒 I ___ a boy. [am✓/is/are] 复习!', '👦 He ___ tall. [am/is✓/are] 复习!', '👧 She ___ happy. [am/is✓/are] 复习!'])
add_tip('题10-12 混合 I am / He is / She is 的复习。这是第一次混合练习，如果错误率突然升高，退回到 Level 1-1 或 1-2 重新练习。')

doc.add_page_break()

# ---- Level 2 ----
doc.add_heading('八、Level 2 详细教案 — 冠词 + 名词复数', level=1)

doc.add_heading('Level 2-1: a vs an', level=2)
add_para('核心规则：a + 辅音开头 / an + 元音开头（a, e, i, o, u）', bold=True, size=10.5)
add_para('ASD 教法：不解释"元音/辅音"术语，而是用"嘴巴张开的声音"。所有 a/an 题目中，单词首字母如果是 a/e/i/o/u 用红色标记，其他字母用灰色标记。', size=10.5)

add_step_box(1, 'SEE（2题）', None)
add_exercise('题1-2:', ['🐱→a cat (c=灰色辅音)', '🍎→an apple (a=红色元音)'])

add_step_box(3, 'MATCH（6题，选 a/an）', None)
add_exercise('题5-10:', ['🐱 ___ cat [a✓/an]', '🍎 ___ apple [a/an✓]', '🐕 ___ dog [a✓/an]', '🥚 ___ egg [a/an✓]', '☂ ___ umbrella [a/an✓]', '📚 ___ book [a✓/an]'])

add_warning('不要让学生背诵"元音字母是 a e i o u"。用视觉（红色字母标记）替代规则记忆。学生看到红色字母就选 an，看到灰色字母就选 a。')

doc.add_heading('Level 2-2: a/an vs the', level=2)
add_para('核心区别：a/an = 泛指的"一个"（随便哪个都可以）；the = 特指的"那个"（就是指这一个）', bold=True, size=10.5)
add_tip('用两张图对比教学：🐱（随便一只猫）vs 🐱👉（手指指着一只特定的猫）。这个视觉对比比任何语言解释都有效。')

doc.add_heading('Level 2-3: 名词复数 -s / -es / -ies', level=2)
add_para('核心规则：一个 → 不加 s；两个及以上 → 加 s/es/ies', bold=True, size=10.5)
add_para('词尾 -s/-es/-ies 用 🟡 黄色标记，延续"词尾变化 = 黄色"的规则。', size=10.5)
add_tip('从本关开始建立"黄色 = 词尾巴有变化"的视觉规则。这个规则会延续到 Level 4（三单-s）和 Level 6（进行时-ing）。')

doc.add_page_break()

# ---- Level 3 ----
doc.add_heading('九、Level 3 详细教案 — 一般现在时（无三单）', level=1)

add_para('Level 3 引入动作动词。注意：刻意排除 He/She/It 作主语（留到 Level 4），本 Level 只用 I/You/We/They + 动词原形。', size=10.5)

doc.add_heading('Level 3-1: I/You/We/They + 动词原形', level=2)
add_para('引入"时间圈"视觉：屏幕左侧固定显示一个圈，里面写着"🌅 每天 every day"，帮助学生理解"一般现在时 = 每天做的事"。', size=10.5)
add_para('动词库：eat breakfast🍳, go to school🏫, play soccer⚽, read books📚, drink water💧, run🏃, sleep😴, sing🎤', size=10)

add_step_box(3, 'MATCH（4题，选动词原形 vs +s）', None)
add_exercise('题5-8:', ['🍳 I ___ breakfast. [eat✓/eats]', '⚽ We ___ soccer. [play✓/plays]', '💧 They ___ water. [drink✓/drinks]', '🏃 I ___ fast. [run✓/runs]'])

doc.add_heading('Level 3-2: 否定句 — don\'t + 动词原形', level=2)
add_para('核心：don\'t = do + not 合并。don\'t 后面的动词永远用原形！', bold=True, size=10.5)
add_tip('视觉设计：don\'t 用紫色（function），配图用摇头😐来传达否定含义。')
add_step_box(3, 'MATCH（4题，肯定 vs 否定）', None)
add_exercise('题5-8:', ['🐟😊 I ___ fish. [like✓/don\'t like]', '🐟😐 I ___ fish. [like/don\'t like✓]', '🏫😐 They ___ to school. [go/don\'t go✓]', '⚽😐 We ___ soccer. [play/don\'t play✓]'])

doc.add_heading('Level 3-3: 疑问句 — Do you...?', level=2)
add_para('核心：Do + 主语 + 动词原形 + ...?', bold=True, size=10.5)
add_tip('视觉设计：Do 从句子外面"跳"到最前面，句号变成问号（动画）。')
add_step_box(3, 'MATCH（4题）', None)
add_exercise('题5-8:', ['🐕❓ ___ you like dogs? [Do✓/Does]', '⚽❓ ___ they play soccer? [Do✓/Does]', '🍳❓ ___ you eat breakfast? [Do✓/Does]', '📚❓ ___ we read books? [Do✓/Does]'])

doc.add_page_break()

# ---- Level 4 ----
doc.add_heading('十、Level 4 详细教案 — 一般现在时（三单）', level=1)

add_para('⚠️ 这是整个体系中最难的一级。ASD 学生需要同时处理：主语识别 + 动词变形 + 发音变化。', bold=True, size=11)
add_para('解决方案：分三步，每步只加一个变量。引入"偷走 -s"的视觉隐喻。', size=10.5)

doc.add_heading('Level 4-1: He/She/It + 动词+s/es（肯定句）', level=2)
add_para('核心：主语是 He/She/It 时，动词要加 -s 或 -es。词尾用 🟡黄色标记。', bold=True, size=10.5)
add_tip('从本关开始建立"he/she/it → 动词+s"的条件反射。不要解释为什么，只要看到 he/she/it 主语，动词就选带 -s 的那个。')

add_step_box('1.5', '🔄 对比观察（2题）', None)
add_exercise('题3-4:', ['🍳🧒 I eat breakfast. → 🍳👦 He eats breakfast. (eat→eats!)', '🏫 I go to school. → 👧 She goes to school. (go→goes!)'])
add_tip('这个对比直接用 Level 3-1 学生已经会的句子，只改变主语，让学生聚焦于动词变化。')

add_step_box(3, 'MATCH（4题）', None)
add_exercise('题7-10:', ['👦🍳 He ___ breakfast. [eat/eats✓]', '👧🏫 She ___ to school. [go/goes✓]', '🐕💧 It ___ water. [drink/drinks✓]', '👦⚽ He ___ soccer. [play/plays✓]'])

doc.add_heading('Level 4-2: doesn\'t 否定', level=2)
add_para('核心规则：doesn\'t = does + not。doesn\'t 后面的动词不加 -s！', bold=True, size=10.5)
add_para('核心比喻：doesn\'t "偷走了"动词的 -s，因为 doesn\'t 里面已经有 does 的三单信息了。', size=10.5)
add_tip('动画设计：肯定句 He eats fish. 中的 -s（黄色）在变成否定句 He doesn\'t eat fish. 时，-s 从 eat 上"飞走"消失。这个视觉动画是 ASD 学生理解 doesn\'t 后动词不加 -s 的关键！')

add_step_box(3, 'MATCH — 关键选择题（4题）', None)
add_exercise('题5-8:', ['👦🐟😊 He ___ fish. [likes✓/doesn\'t like]', '👦🐟😐 He ___ fish. [likes/doesn\'t like✓]',
    '👧⚽😐 She ___ soccer. [doesn\'t play✓/doesn\'t plays] ←关键！', '👦💧😐 He ___ water. [doesn\'t drink✓/doesn\'t drinks] ←关键！'])
add_warning('题7和题8是理解 doesn\'t 的核心！选项 "doesn\'t plays" 是故意设置的干扰项。如果学生选了这个，说明他还没理解"doesn\'t 偷走 -s"。这时候需要回到对比观察环节再看一次动画。')

doc.add_heading('Level 4-3: Does...? 疑问', level=2)
add_para('核心规则：Does + 主语 + 动词原形...?（Does 后面的动词也不加 -s！）', bold=True, size=10.5)
add_para('同样的比喻：Does 也"偷走了"动词的 -s。', size=10.5)

add_step_box(3, 'MATCH', None)
add_exercise('题5-8:', ['👦🐕❓ ___ he like dogs? [Do/Does✓]', '👧⚽❓ ___ she play soccer? [Do/Does✓]', '👦🍳❓ ___ he eat breakfast? [Do/Does✓]', '🐕💧❓ ___ it drink water? [Do/Does✓]'])

doc.add_page_break()

# ---- Level 5 ----
doc.add_heading('十一、Level 5 详细教案 — 情态动词 Can', level=1)

add_para('为什么 can/can\'t 放在三单之后？因为 can 后面的动词永远不变！在三单的艰难学习后，can 给学生一个"喘息"的机会，重建信心。', size=10.5)

add_para('核心规则（极简）：can 后面的动词永远用原形，不管主语是谁！', bold=True, size=10.5)
add_tip('不需要说"情态动词"这个词。只需要说"can 是一个小帮手词（紫色），它后面的动作词永远不变。"')

doc.add_heading('Level 5-1: I can / I can\'t', level=2)
add_para('能力动词：swim🏊, run fast🏃💨, jump🤾, read📖, cook🍳, sing🎤, dance💃, fly🕊️', size=10)
add_exercise('MATCH:', ['🏊 I ___ swim. [can✓/can\'t]', '🏊❌ I ___ swim. [can/can\'t✓]', '🕊️ Birds ___ fly. [can✓/can\'t]', '🐕🕊️ Dogs ___ fly. [can/can\'t✓]'])
add_tip('can\'t 的配图使用 ✅/❌ 而不是面部表情，因为 ASD 学生可能难以解读表情。')

doc.add_heading('Level 5-2: Can you...? 问答', level=2)
add_para('核心：Can you + 动词原形...? → Yes, I can. / No, I can\'t.', bold=True, size=10.5)

doc.add_heading('Level 5-3: He/She can', level=2)
add_para('关键：He/She can swim.（不是 cans！can 永远不变）', bold=True, size=10.5)
add_tip('这关很容易——学生已经通过 L4 习惯了 he/she 后面动词要变，但 can 是例外。出 1-2 题 "He cans swim" 作为干扰项，让学生自己发现 can 不变。')

doc.add_page_break()

# ---- Level 6 ----
doc.add_heading('十二、Level 6 详细教案 — 现在进行时', level=1)

add_para('现在进行时 = L1（be动词） + L3（动作动词） + -ing。学生已经有了前两个基础，现在只需要把 -ing 加上去。', size=10.5)

add_para('核心公式：主语 + am/is/are + 动词-ing', bold=True, size=10.5)
add_para('视觉标记：-ing 用 🟡黄色（与三单 -s 同色系，延续"词尾变化 = 黄色"的规则）', size=10.5)

doc.add_heading('Level 6-1: I am + verb-ing', level=2)
add_para('-ing 拼写规则分三小步：6-1 一般动词 +ing（eat→eating）；6-2 去 e +ing（make→making）；6-3 双写 +ing（run→running）。每小步只在对应题目中自然出现，不集中讲解规则。', size=10)
add_step_box(3, 'MATCH', None)
add_exercise('题5-8:', ['📖 I ___ reading. [am✓/is]', '😴 I am ___. [sleep/sleeping✓]', '⚽ I am ___ soccer. [play/playing✓]', '🍳 I ___ eating. [am✓/is]'])

doc.add_heading('Level 6-2: He/She is + verb-ing', level=2)
add_para('核心区别：am → is（主语变化）', bold=True, size=10.5)

doc.add_heading('Level 6-3: They are + verb-ing', level=2)
add_para('核心区别：is → are（复数主语）', bold=True, size=10.5)

doc.add_page_break()

# ---- Level 7 ----
doc.add_heading('十三、Level 7 详细教案 — 介词方位', level=1)

add_para('这是 ASD 学生最有优势的一个 Level！空间关系是具象的、可视觉化的，许多 ASD 学生在空间思维上表现优异。', size=10.5)
add_tip('所有题目使用同一张桌子图片 + 一只小猫在不同位置。利用 ASD 学生的模式识别优势。介词统一用 🟣紫色标记。')

doc.add_heading('Level 7-1: in / on / under', level=2)
add_para('三个最基础的空间介词：in（在里面）、on（在上面）、under（在下面）', size=10.5)
add_exercise('MATCH:', ['🐱📦 The cat is ___ the box. [in✓/on/under]', '🐱📦 The cat is ___ the box. [in/on✓/under]', '🐱📦 The cat is ___ the box. [in/on/under✓]'])
add_tip('三张图并排展示：猫在盒子里🐱📦 / 猫在盒子上🐱📦 / 猫在盒子下🐱📦。位置对比一目了然。')

doc.add_heading('Level 7-2: next to / behind / in front of', level=2)
add_para('更复杂的空间关系，需要更多参照点判断。', size=10.5)

doc.add_heading('Level 7-3: There is/are + 介词', level=2)
add_para('引入存在句：There is a book on the desk. / There are two cats under the box.', bold=True, size=10.5)
add_para('关键区别：单数用 There is，复数用 There are。', size=10.5)

doc.add_page_break()

# ---- Level 8 ----
doc.add_heading('十四、Level 8 详细教案 — Have / Has', level=1)

add_para('have/has 的规则与 am/is/are 非常相似。学生已经在 Level 1 掌握了人称-Be动词的对应关系，现在只需要做类比迁移。', size=10.5)

add_para('类比表（核心教学工具）：', bold=True, size=10.5)
add_table(['Be 动词（已学 ✓）', '→', 'Have/Has（新学 🆕）'], [
    ['I am', '→', 'I have'],
    ['He is', '→', 'He has  ← 都有 s！'],
    ['She is', '→', 'She has  ← 都有 s！'],
    ['You are', '→', 'You have'],
    ['We are', '→', 'We have'],
    ['They are', '→', 'They have'],
], [4.5, 1.0, 7.5])

add_tip('关键类比：is 和 has 都有 s。只要主语用 he/she/it，动词就选带 s 的那个。学生不需要理解"为什么"，只需要建立模式识别。')

doc.add_heading('Level 8-3: 综合复习', level=2)
add_para('混合 am/is/are、have/has、can/can\'t、现在进行时、介词、三单。', size=10.5)
add_para('全部为 MATCH + BUILD + USE，不再有 SEE/COLOR。检验学生是否能在无提示情况下正确选择。', size=10.5)
add_warning('综合复习是检验泛化能力的关键。如果学生在混合练习中错误率超过 30%，说明之前的模式是在"特定关卡"中建立的，还没有真正泛化。需要回到出错最多的关卡重新练习。')

doc.add_page_break()

# ====================================================================
# 十五、情绪安全反馈规范
# ====================================================================
doc.add_heading('十五、情绪安全反馈规范', level=1)

add_para('ASD 学生的高度焦虑意味着反馈系统必须精心设计。以下规范适用于所有关卡的所有题型。', size=10.5)

doc.add_heading('15.1 正确反馈', level=2)
add_table(['反馈方式', '何时使用', '说明'], [
    ['⭐ 星星 +1', '每次正确', '屏幕右上角星星数 +1，有轻微缩放动画'],
    ['🟢 绿色边框', '选项按钮/输入框正确', '边框变为绿色，持续 1 秒'],
    ['"Good job!"', '连续 3 题正确', '短暂文字提示，2 秒后消失'],
    ['🌟 庆祝动画', '每 10 颗星', '星星从屏幕中央飞散（无声音，无弹窗）'],
    ['🏆 关卡完成', '通关时', '奖杯图标 + 获得的星星数 + 鼓励文字'],
], [3.0, 3.5, 8.5])

doc.add_heading('15.2 错误处理（最重要）', level=2)
add_table(['❌ 绝对禁止', '原因', '✅ 替代做法'], [
    ['红色边框/红色文字', '引发焦虑，红色=危险', '黄色闪烁 1.5 秒'],
    ['"错了！""不对！"', '批评性语言导致退缩', '不评价，直接显示答案'],
    ['❌ 叉号图标', '负面视觉刺激', '💡 提示图标'],
    ['分数/排名', '社交比较焦虑', '只有 ⭐ 星星累计'],
    ['"再想想"', '给 ASD 学生施加认知压力', '自动显示答案，然后下一题'],
    ['错误音效', '感官刺激 + 负面条件反射', '无音效，只有视觉过渡'],
], [3.5, 4.5, 7.0])

doc.add_heading('15.3 错误处理流程', level=2)
add_para('1. 学生选了错误答案 → 错误选项短暂黄色闪烁（1.5 秒）', size=10.5)
add_para('2. 正确答案自动以绿色显示（不解释为什么不选那个）', size=10.5)
add_para('3. 无任何文字评价', size=10.5)
add_para('4. 2-3 秒后自动进入下一题', size=10.5)
add_para('5. 该题的正确答案会在 1 天后出现在"间隔复习"中', size=10.5)

doc.add_page_break()

# ====================================================================
# 十六、教师使用指南
# ====================================================================
doc.add_heading('十六、教师使用指南', level=1)

doc.add_heading('16.1 如何使用本教案', level=2)

add_para('1. 按顺序推进：Level 1 → Level 2 → ... → Level 8，不要跳级。每个 Level 内部的 3 关也要按顺序。', size=10.5)
add_para('2. 过关标准：当前关 Step 5（USE）正确率 ≥ 75% 才能解锁下一关。不要在学生还没准备好时强行推进。', size=10.5)
add_para('3. 每关时长控制：8-15 分钟。如果学生在 8 分钟内完成不了，说明难度过高，降低提示等级。如果超过 15 分钟还没完成，允许暂停，下次继续。', size=10.5)
add_para('4. 观察而非纠正：学生在 MATCH 阶段连续选错时，不要说"应该是 is"，而是等待系统自动显示答案。过度纠正会引发焦虑和抗拒。', size=10.5)
add_para('5. 回退是正常的：如果学生在 Level 4（三单）卡住超过 3 次尝试，建议退回到 Level 3（一般现在时）重新练习。回退不是失败，是巩固。', size=10.5)

doc.add_heading('16.2 ASD 学生常见反应及应对', level=2)

add_table(['学生表现', '可能原因', '教师应对'], [
    ['反复做同一关不换', '喜欢可预测性/刻板行为', '温和引导到下一关，说"你已经很棒了，我们试试新的"'],
    ['拒绝开口跟读', '感官敏感/发音焦虑', '不强制。SEE 环节的 TTS 是被动听的，不需要跟读'],
    ['看到颜色标记反而困惑', '视觉过度刺激', '手动调低提示等级到 L2，减少颜色'],
    ['连续快速点击不思考', '想快点结束/逃避', '设置每题最低停留时间（3秒），但不能用计时器'],
    ['在 USE 写句阶段完全空白', '泛化能力不足', '回退到 BUILD，再练几题，然后打开💡提示按钮'],
    ['答对了也不开心', '面部表情解读困难/情绪表达方式不同', '不需要强行调动情绪。星星累计本身就是正向反馈'],
], [3.5, 4.0, 7.5])

doc.add_heading('16.3 学习节奏建议', level=2)

add_para('• 每天 1-2 关（15-30 分钟），不要连续做超过 2 关', size=10.5)
add_para('• 每完成 3 关（=1 个 Level），安排 1 天休息，让知识"沉淀"', size=10.5)
add_para('• 新 Level 开始前，先复习上一个 Level 的 USE 题（2-3 题即可）', size=10.5)
add_para('• 预计总时间：24 关 × 10 分钟/关 = 约 4 小时（不含休息）。按每天 2 关计算，约 12 个学习日完成全部', size=10.5)

doc.add_heading('16.4 评估追踪', level=2)
add_para('不需要考试或测验来评估。关注以下指标：', size=10.5)
add_para('• 每题平均耗时（突然增长 = 卡点）', size=10.5)
add_para('• 渐隐等级变化（向下趋势 = 在进步）', size=10.5)
add_para('• 同类型错误重复次数（>3 次 = 需要回退）', size=10.5)
add_para('• USE 阶段正确率（< 50% = 泛化不足，需加强 MATCH/BUILD）', size=10.5)

# ====================================================================
# 尾页
# ====================================================================
doc.add_page_break()
doc.add_heading('📋 附录：词汇范围总表', level=1)

vocab = [
    ['人称代词', '7', 'I, you, he, she, it, we, they'],
    ['日常动词', '20', 'eat, drink, sleep, run, walk, read, write, play, sing, dance, swim, jump, cook, watch, like, love, go, come, sit, stand'],
    ['形容词', '12', 'happy, sad, tall, short, big, small, hungry, thirsty, tired, hot, cold, good'],
    ['身份名词', '6', 'boy, girl, student, teacher, friend, baby'],
    ['物品名词', '15', 'book, pen, desk, chair, table, box, bag, ball, cat, dog, fish, bird, apple, egg, umbrella'],
    ['食物', '8', 'apple, egg, rice, bread, milk, water, fish, cake'],
    ['地点', '5', 'school, home, park, shop, room'],
    ['介词', '6', 'in, on, under, next to, behind, in front of'],
]
add_table(['类别', '数量', '词汇'], vocab, [3.0, 1.5, 10.5])

add_para('')
add_para('— 文档结束 —', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))
add_para(f'生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}', size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))
add_para('本教案专为 ASD 青少年设计，遵循"循序渐进、视觉引导、情绪安全"三原则。', size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))

# ============ 保存 ============
output_path = r'C:\西安备课\语法训练\ASD语法训练教案.docx'
doc.save(output_path)
print(f'Done: {output_path}')
# Also copy to desktop
import shutil
desktop = r'C:\Users\25152\Desktop\ASD语法训练教案.docx'
shutil.copy2(output_path, desktop)
print(f'Copied to: {desktop}')
